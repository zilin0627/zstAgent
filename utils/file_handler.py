import hashlib
import os
from typing import Any

import pymupdf
from docx import Document as DocxDocument
from langchain_core.documents import Document
from langchain_community.document_loaders import PyPDFLoader, TextLoader
from pypdf import PdfReader

from utils.logger_handler import logger


def get_file_md5_hex(file_path: str) -> str:
    """
    get_file_md5_hex 的 Docstring
    传入文件路径，返回文件的MD5哈希值的十六进制表示
    """
    if not os.path.exists(file_path):
        logger.error(f"[md5计算]文件不存在: {file_path}")
        return None

    if not os.path.isfile(file_path):
        logger.error(f"[md5计算]路径不是文件: {file_path}")
        return None

    md5_obj = hashlib.md5()
    chunk_size = 4096  # 4KB 读取文件, 避免内存溢出
    try:
        with open(file_path, "rb") as f:
            while chunk := f.read(chunk_size):  # 按块读取大文件
                md5_obj.update(chunk)  # 增量更新MD5哈希值
            return md5_obj.hexdigest()
    except Exception as e:
        logger.error(f"[md5计算]读取文件 {file_path} 时出错: {str(e)}")
        return None


def listdir_with_allowed_type(dir_path: str, allowed_types: tuple[str]):
    """
    列出目录下所有符合类型要求的文件
    传入目录路径和允许的文件类型元组，返回符合要求的文件列表
    """
    allowed_files = []
    if not os.path.isdir(dir_path):
        logger.error(f"[listdir_with_allowed_type]路径不是目录: {dir_path}")
        return []

    for root, _, files in os.walk(dir_path):
        for file in files:
            if file.endswith(allowed_types):  # 检查文件是否以允许的类型结尾
                allowed_files.append(os.path.join(root, file))

    return allowed_files


def _normalize_page_number(value: Any) -> int | None:
    if isinstance(value, int):
        return value
    if isinstance(value, str) and value.isdigit():
        return int(value)
    return None


def _decorate_documents(
    docs: list[Document],
    *,
    source: str,
    parser_name: str,
    parser_strategy: str = "",
    has_ocr: bool = False,
) -> list[Document]:
    decorated_docs: list[Document] = []
    for page_index, doc in enumerate(docs, start=1):
        metadata = dict(doc.metadata or {})
        page_number = _normalize_page_number(
            metadata.get("page_number", metadata.get("page", page_index))
        )
        element_type = str(
            metadata.get("category")
            or metadata.get("element_type")
            or metadata.get("type")
            or "text"
        )
        text_as_html = metadata.get("text_as_html") or metadata.get("table_as_html") or ""
        metadata.update(
            {
                "source": source,
                "page": page_number if page_number is not None else page_index,
                "page_number": page_number if page_number is not None else page_index,
                "element_type": element_type,
                "source_parser": parser_name,
                "parser_strategy": parser_strategy,
                "has_ocr": has_ocr,
                "has_table": element_type.lower() == "table",
                "table_html": text_as_html,
                "has_image": element_type.lower() in {"image", "figurecaption"},
            }
        )
        decorated_docs.append(Document(page_content=doc.page_content, metadata=metadata))
    return decorated_docs


def _pypdf_loader(file_path: str, password: str = None) -> list[Document]:
    docs = PyPDFLoader(file_path, password=password).load()
    return _decorate_documents(docs, source=file_path, parser_name="pypdfloader")


def _fallback_pdf_reader(file_path: str, password: str = None) -> list[Document]:
    reader = PdfReader(file_path)
    if reader.is_encrypted:
        try:
            decrypt_result = reader.decrypt(password or "")
            if decrypt_result == 0:
                logger.warning(f"[pdf_loader]PDF处于加密状态但空密码无法解开，继续尝试直接读取: {file_path}")
        except Exception as decrypt_error:
            logger.warning(f"[pdf_loader]PDF解密尝试失败，继续尝试直接读取: {file_path} | {str(decrypt_error)}")

    docs: list[Document] = []
    for page_index, page in enumerate(reader.pages, start=1):
        try:
            text = page.extract_text() or ""
        except Exception as page_error:
            logger.warning(f"[pdf_loader]PDF第 {page_index} 页提取失败，跳过: {file_path} | {str(page_error)}")
            text = ""
        if not text.strip():
            continue
        docs.append(Document(page_content=text, metadata={"source": file_path, "page": page_index}))
    return _decorate_documents(docs, source=file_path, parser_name="pypdf")


def _fallback_pymupdf_reader(file_path: str) -> list[Document]:
    pdf = pymupdf.open(file_path)
    docs: list[Document] = []
    try:
        if pdf.needs_pass:
            auth_ok = pdf.authenticate("")
            if not auth_ok:
                logger.warning(f"[pdf_loader]PyMuPDF空密码认证失败，继续尝试直接读取: {file_path}")

        for page_index in range(pdf.page_count):
            try:
                page = pdf.load_page(page_index)
                text = page.get_text("text") or ""
            except Exception as page_error:
                logger.warning(f"[pdf_loader]PyMuPDF第 {page_index + 1} 页提取失败，跳过: {file_path} | {str(page_error)}")
                text = ""
            if not text.strip():
                continue
            docs.append(
                Document(
                    page_content=text,
                    metadata={"source": file_path, "page": page_index + 1},
                )
            )
    finally:
        pdf.close()
    return _decorate_documents(docs, source=file_path, parser_name="pymupdf")


def pdf_loader(file_path: str, password: str = None) -> list[Document]:
    """
    加载PDF文件
    传入PDF文件路径，返回PDF文件的内容列表
    """
    try:
        docs = _pypdf_loader(file_path, password=password)
        if docs:
            logger.info(f"[pdf_loader]PyPDFLoader读取成功: {file_path}，共 {len(docs)} 页")
            return docs
    except Exception as e:
        logger.warning(f"[pdf_loader]PyPDFLoader加载失败，尝试继续回退读取: {file_path} | {str(e)}")

    try:
        docs = _fallback_pdf_reader(file_path, password=password)
        if docs:
            logger.info(f"[pdf_loader]pypdf回退读取成功: {file_path}，共 {len(docs)} 页")
            return docs
    except Exception as fallback_error:
        logger.warning(f"[pdf_loader]pypdf回退读取失败，继续尝试PyMuPDF: {file_path} | {str(fallback_error)}")

    try:
        docs = _fallback_pymupdf_reader(file_path)
        if docs:
            logger.info(f"[pdf_loader]PyMuPDF回退读取成功: {file_path}，共 {len(docs)} 页")
            return docs
        logger.error(f"[pdf_loader]回退读取后仍无可用内容: {file_path}")
        return None
    except Exception as pymupdf_error:
        logger.error(f"[pdf_loader]加载PDF文件 {file_path} 时出错: {str(pymupdf_error)}")
        return None


def txt_loader(file_path: str) -> list[Document]:
    """
    加载TXT文件
    传入TXT文件路径，返回TXT文件的内容列表
    """
    try:
        docs = TextLoader(file_path, encoding="utf-8").load()
        return _decorate_documents(docs, source=file_path, parser_name="text")
    except Exception as e:
        logger.error(f"[txt_loader]加载TXT文件 {file_path} 时出错: {str(e)}")
        return None


def docx_loader(file_path: str) -> list[Document]:
    try:
        doc = DocxDocument(file_path)
        text = "\n".join(p.text for p in doc.paragraphs if p.text and p.text.strip())
        if not text.strip():
            return []
        docs = [Document(page_content=text, metadata={"source": file_path, "page": 1})]
        return _decorate_documents(docs, source=file_path, parser_name="docx")
    except Exception as e:
        logger.error(f"[docx_loader]加载DOCX文件 {file_path} 时出错: {str(e)}")
        return None


def load_with_txt_fallback(file_path: str, password: str = None) -> tuple[list[Document], str]:
    if file_path.endswith(".txt"):
        docs = txt_loader(file_path) or []
        return docs, file_path if docs else ""

    if file_path.endswith(".docx"):
        docs = docx_loader(file_path) or []
        return docs, file_path if docs else ""

    if file_path.endswith(".pdf"):
        docs = pdf_loader(file_path, password=password) or []
        if docs:
            return docs, file_path

        txt_path = os.path.splitext(file_path)[0] + ".txt"
        if os.path.exists(txt_path):
            logger.warning(f"[pdf_loader]PDF读取失败，使用同名TXT兜底: {txt_path}")
            txt_docs = txt_loader(txt_path) or []
            if txt_docs:
                return txt_docs, txt_path

        docx_path = os.path.splitext(file_path)[0] + ".docx"
        if os.path.exists(docx_path):
            logger.warning(f"[pdf_loader]PDF读取失败，使用同名DOCX兜底: {docx_path}")
            docx_docs = docx_loader(docx_path) or []
            if docx_docs:
                return docx_docs, docx_path
        return [], ""

    return [], ""
